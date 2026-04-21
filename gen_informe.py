#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Genera docs/Informe_ProyectoFinal_COMP2315.docx
Sistema de Expedientes Médicos del Dr. Rodríguez — COMP 2315

Formato:
  - Times New Roman 12 pt
  - Interlineado 1.5
  - Márgenes 1 pulgada (todos los lados)
  - Números de página (encabezado superior derecho)
  - Portada, TOC, Introducción, Bitácora, Flujogramas,
    Menú, Código fuente (6 archivos), Conclusión, Referencias
"""

import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION_START, WD_ORIENTATION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy
import lxml.etree as etree

# ── Rutas ──────────────────────────────────────────────────────────────────────
JAVA_DIR = os.path.join(os.path.dirname(__file__), 'java')
OUT_PATH = os.path.join(os.path.dirname(__file__), 'docs',
                        'Informe_ProyectoFinal_COMP2315.docx')

# ── Datos del proyecto ─────────────────────────────────────────────────────────
TEAM = [
    ('Ektor M. Gonzalez Ramos',       'A00617167'),
    ('Diego L. Rodriguez Perez',       'A00636906'),
    ('Donovan Irizarry Figueroa',      'A00671799'),
    ('Carlo Ramos',                   'A00642991'),
]
COURSE    = 'COMP 2315 — Programación Estructurada'
SECTION   = '202630.73752'
PROFESOR  = 'Dr. Edgardo Vargas Moya'
INST      = 'Universidad Interamericana de Puerto Rico\nRecinto de Aguadilla\nFacultad de Ciencias y Tecnología'
YEAR      = '2026'
TITLE     = 'Sistema de Expedientes Médicos del Dr. Rodríguez'
SUBTITLE  = 'Proyecto Final'

JAVA_FILES = [
    ('Main.java',             'Punto de entrada del sistema. Coordina bienvenida, login, menú principal con switch y cierre.'),
    ('Paciente.java',         'Modelo del expediente médico. Contiene 13 campos, getters, setters y serialización a archivo.'),
    ('Validador.java',        'Capa de validación. Verifica formatos con charAt() y length() (SS, fechas, expediente).'),
    ('ArchivoManager.java',   'Persistencia en archivo. Guarda, carga y reescribe expedientes en expedientes.txt.'),
    ('ExpedienteManager.java','Lógica CRUD. Crea, busca por número/fecha, actualiza y lista expedientes.'),
    ('MenuUtils.java',        'Interfaz de usuario. Bienvenida, despedida, login con 3 intentos y menú principal via JOptionPane.'),
]

BITACORA = [
    ('04/06/2026', '6:00 PM – 8:00 PM\n(2 horas)',
     'Todos los integrantes',
     'Análisis de requisitos del proyecto y diseño de clases; lectura de rúbrica y distribución de responsabilidades.',
     'Diagrama conceptual de las seis clases; definición de los 13 campos del expediente; configuración del proyecto en el IDE.'),
    ('04/08/2026', '7:00 PM – 9:00 PM\n(2 horas)',
     'Ektor M. Gonzalez\nDiego L. Rodriguez',
     'Implementación de Paciente.java y Validador.java: atributos, constructor, getters, setters y serialización.',
     'Clase Paciente completamente funcional con toFileString() / fromFileString(); clase Validador con validaciones usando charAt() y length().'),
    ('04/10/2026', '6:30 PM – 9:30 PM\n(3 horas)',
     'Donovan Irizarry\nCarlo Ramos',
     'Implementación de ArchivoManager.java y ExpedienteManager.java: lógica de archivos y operaciones CRUD.',
     'Persistencia con FileReader/PrintWriter; búsqueda por número y por fecha de visita; submenú de actualización con switch de 11 opciones.'),
    ('04/12/2026', '5:00 PM – 7:30 PM\n(2.5 horas)',
     'Todos los integrantes',
     'Implementación de MenuUtils.java y Main.java: login con control de intentos, menú principal con JOptionPane, integración de módulos.',
     'Sistema de autenticación con arreglos paralelos (USUARIOS/CONTRASENAS), límite de 3 intentos; menú de 6 opciones; bucle while del flujo principal.'),
    ('04/14/2026', '4:00 PM – 6:00 PM\n(2 horas)',
     'Todos los integrantes',
     'Pruebas de integración, corrección de errores, revisión de documentación y preparación del informe y la presentación.',
     'Sistema probado con casos de prueba completos (campos vacíos, formatos incorrectos, cancelaciones); todos los errores críticos corregidos; documentación finalizada.'),
]

INTRO = (
    'El presente documento constituye el informe final del proyecto de programación '
    'desarrollado para el curso COMP 2315 — Programación Estructurada, bajo la '
    'dirección del Dr. Edgardo Vargas Moya en la Universidad Interamericana de '
    'Puerto Rico, Recinto de Aguadilla. El proyecto consiste en el diseño e '
    'implementación de un Sistema de Expedientes Médicos del Dr. Rodríguez, una '
    'aplicación de consola e interfaz gráfica básica desarrollada íntegramente en '
    'Java, que permite gestionar de forma digital los registros clínicos de los '
    'pacientes de un consultorio médico privado.',

    'El sistema fue construido con el propósito de aplicar los conocimientos '
    'adquiridos a lo largo del semestre. Entre sus funcionalidades principales se '
    'encuentran: un módulo de autenticación con control de intentos (login), '
    'operaciones CRUD (crear, buscar, actualizar y listar expedientes), persistencia '
    'de datos en archivos de texto plano, y validación de datos usando charAt() y '
    'length() para garantizar la integridad de la información. La interfaz está '
    'implementada mediante JOptionPane de la biblioteca javax.swing. La arquitectura '
    'se organiza en seis clases: Main, Paciente, Validador, ArchivoManager, '
    'ExpedienteManager y MenuUtils, aplicando principios de diseño visto en el curso.',
)

CONCLUSION = (
    'El desarrollo del Sistema de Expedientes Médicos del Dr. Rodríguez nos '
    'permitió aplicar de manera integrada todos los conceptos estudiados en '
    'COMP 2315 — Programación Estructurada. Escogimos este proyecto porque '
    'representa una necesidad real en el entorno de la salud: la digitalización '
    'y gestión eficiente de expedientes médicos, lo que nos motivó a crear una '
    'solución funcional y robusta.',

    'A lo largo del desarrollo experimentamos retos significativos, entre ellos: '
    'el diseño de la validación de formato sin el uso de expresiones regulares '
    '(solo charAt() y length() para cumplir el currículo), la gestión del '
    'arreglo fijo Paciente[100] evitando desbordamientos, y la coordinación '
    'del trabajo en equipo para integrar los seis módulos sin conflictos. '
    'Cada uno de estos retos fortaleció nuestro entendimiento de la programación '
    'estructurada y del trabajo colaborativo en un proyecto de software real.',

    'Como conclusión, este proyecto nos demostró que con las estructuras básicas '
    '— arreglos, métodos, if/else, while y switch — es posible construir una '
    'aplicación completa y útil. La separación en clases facilitó el '
    'mantenimiento y la depuración del código, y el uso de archivos de texto '
    'plano demostró la importancia de la persistencia de datos incluso sin '
    'bases de datos relacionales.',
)

REFERENCIAS = [
    'Oracle Corporation. (2024). Java SE Documentation. '
    'https://docs.oracle.com/en/java/javase/',

    'Deitel, P., & Deitel, H. (2020). Java: How to Program (11.ª ed.). Pearson.',

    'Oracle Corporation. (2024). Class JOptionPane. Java Platform SE. '
    'https://docs.oracle.com/javase/8/docs/api/javax/swing/JOptionPane.html',

    'Oracle Corporation. (2024). Class Scanner. Java Platform SE. '
    'https://docs.oracle.com/javase/8/docs/api/java/util/Scanner.html',

    'Oracle Corporation. (2024). Class PrintWriter. Java Platform SE. '
    'https://docs.oracle.com/javase/8/docs/api/java/io/PrintWriter.html',
]

MENU_PRINCIPAL_TEXT = (
    'El menú principal del sistema se presenta al usuario mediante un cuadro de '
    'diálogo JOptionPane.showInputDialog() con seis opciones numeradas. El usuario '
    'ingresa el número de la opción deseada y el sistema ejecuta la acción '
    'correspondiente mediante una estructura switch. Si el usuario cierra el '
    'diálogo sin seleccionar, el sistema trata esa acción como "Salir" (opción 6). '
    'El bucle principal usa una variable booleana "ejecutando" que solo se pone '
    'en false cuando el usuario confirma la salida escribiendo "S".'
)

MENU_OPTIONS = [
    ('1', 'Crear nuevo expediente',   'ExpedienteManager.crearExpediente()'),
    ('2', 'Buscar por número',        'ExpedienteManager.buscarPorNumero()'),
    ('3', 'Buscar por fecha de visita','ExpedienteManager.buscarPorFechaVisita()'),
    ('4', 'Actualizar expediente',    'ExpedienteManager.actualizarExpediente()'),
    ('5', 'Listar todos los expedientes','ExpedienteManager.listarTodosLosExpedientes()'),
    ('6', 'Salir',                    'Solicita confirmación con showInputDialog; si resp=="S" → ejecutando=false'),
]

SUBMENU_TEXT = (
    'El submenú de actualización se presenta cuando el usuario elige la opción 4. '
    'Primero se solicita el número de expediente a modificar. Una vez encontrado, '
    'se entra en un ciclo while (continuar = true) donde se presenta un nuevo '
    'cuadro de diálogo con las 11 opciones de campos a actualizar. El usuario '
    'puede modificar uno o varios campos en la misma sesión antes de terminar '
    'con la opción 11.'
)

SUBMENU_OPTIONS = [
    ('1',  'Nombre completo'),
    ('2',  'Seguro Social (XXX-XX-XXXX)'),
    ('3',  'Fecha de Nacimiento (MM/DD/YYYY)'),
    ('4',  'Sexo (M / F)'),
    ('5',  'Dirección'),
    ('6',  'Plan Médico'),
    ('7',  'Fecha de Visita (MM/DD/YYYY)'),
    ('8',  'Diagnóstico'),
    ('9',  'Receta'),
    ('10', 'Fecha de Siguiente Visita (MM/DD/YYYY)'),
    ('11', 'Terminar actualización'),
]

# ── Helpers de formato ─────────────────────────────────────────────────────────

def set_font(run, name='Times New Roman', size=12, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def set_para_format(para, align=WD_ALIGN_PARAGRAPH.LEFT, first_line=0,
                    space_before=0, space_after=6, keep_with_next=False):
    pf = para.paragraph_format
    pf.alignment = align
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.first_line_indent = Inches(first_line) if first_line else None
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.keep_with_next = keep_with_next

def add_body_para(doc, text, indent=True, align=WD_ALIGN_PARAGRAPH.LEFT,
                  space_before=0, space_after=6):
    p = doc.add_paragraph()
    set_para_format(p, align=align, first_line=0.5 if indent else 0,
                    space_before=space_before, space_after=space_after)
    run = p.add_run(text)
    set_font(run)
    return p

def add_heading1(doc, text):
    p = doc.add_paragraph()
    set_para_format(p, align=WD_ALIGN_PARAGRAPH.CENTER,
                    space_before=12, space_after=6, keep_with_next=True)
    run = p.add_run(text.upper())
    set_font(run, bold=True, size=14)
    # Underline via border bottom on paragraph
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def add_heading2(doc, text):
    p = doc.add_paragraph()
    set_para_format(p, align=WD_ALIGN_PARAGRAPH.LEFT,
                    space_before=10, space_after=4, keep_with_next=True)
    run = p.add_run(text)
    set_font(run, bold=True, size=13)
    return p

def add_heading3(doc, text):
    p = doc.add_paragraph()
    set_para_format(p, align=WD_ALIGN_PARAGRAPH.LEFT,
                    space_before=8, space_after=3, keep_with_next=True)
    run = p.add_run(text)
    set_font(run, bold=True, size=12)
    return p

def add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(__import__('docx.enum.text',
                                      fromlist=['WD_BREAK']).WD_BREAK.PAGE)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

def add_monospace(doc, text, size=9):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.first_line_indent = None
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(size)
    return p

def set_cell_font(cell, text, bold=False, size=11, bg=None):
    """Clear cell and set text with font."""
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.font.bold = bold
    if bg:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), bg)
        tcPr.append(shd)

def add_page_number_header(doc):
    """Add page number to the right side of the header."""
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    p.clear()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run()
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    # Insert page number field
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

def set_section_margins(section, top=1, bottom=1, left=1, right=1):
    section.top_margin    = Inches(top)
    section.bottom_margin = Inches(bottom)
    section.left_margin   = Inches(left)
    section.right_margin  = Inches(right)
    section.header_distance = Inches(0.4)

# ── Portada (sin número de página) ────────────────────────────────────────────

def build_portada(doc):
    # Institución
    p = doc.add_paragraph()
    set_para_format(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=72, space_after=12)
    run = p.add_run(INST)
    set_font(run, bold=True, size=14)

    # Línea divisora
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1'); bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom); pPr.append(pBdr)
    set_para_format(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=24)

    # Título principal
    p = doc.add_paragraph()
    set_para_format(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=6)
    run = p.add_run(TITLE)
    set_font(run, bold=True, size=18)

    # Subtítulo
    p = doc.add_paragraph()
    set_para_format(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=30)
    run = p.add_run(SUBTITLE)
    set_font(run, bold=True, size=14)

    # Línea divisora
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1'); bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom); pPr.append(pBdr)
    set_para_format(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=24)

    # Integrantes
    for nombre, num in TEAM:
        p = doc.add_paragraph()
        set_para_format(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=2)
        r1 = p.add_run(nombre)
        set_font(r1, bold=True)
        r2 = p.add_run(f'  —  {num}')
        set_font(r2)

    # Datos del curso
    datos = [
        ('Curso:', COURSE),
        ('Sección:', SECTION),
        ('Profesor:', PROFESOR),
        ('Fecha de Entrega:', 'Por anunciarse'),
    ]
    for label, value in datos:
        p = doc.add_paragraph()
        set_para_format(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=4, space_after=2)
        r1 = p.add_run(f'{label} ')
        set_font(r1, bold=True)
        r2 = p.add_run(value)
        set_font(r2)

    # Línea divisora
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1'); bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom); pPr.append(pBdr)
    set_para_format(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=24, space_after=6)

    # Año
    p = doc.add_paragraph()
    set_para_format(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=0)
    run = p.add_run(YEAR)
    set_font(run, bold=True, size=14)

# ── Tabla de Contenido ─────────────────────────────────────────────────────────

TOC_ENTRIES = [
    (0, '1. Introducción',                              '3'),
    (0, '2. Diseño',                                    '4'),
    (1, '2.1 Bitácora de Trabajo',                      '4'),
    (1, '2.2 Flujogramas',                              '5'),
    (2, '2.2.1 Flujo Principal (Main.java)',             '5'),
    (2, '2.2.2 Proceso de Login (MenuUtils.java)',       '5'),
    (2, '2.2.3 Crear Expediente (ExpedienteManager)',    '5'),
    (2, '2.2.4 Buscar por Número (ExpedienteManager)',   '5'),
    (2, '2.2.5 Buscar por Fecha (ExpedienteManager)',    '6'),
    (2, '2.2.6 Actualizar Expediente (ExpedienteManager)','6'),
    (2, '2.2.7 Listar Expedientes (ExpedienteManager)',  '6'),
    (1, '2.3 Menú Principal y Submenú de Actualización', '7'),
    (1, '2.4 Código Fuente',                             '8'),
    (2, 'Main.java',                                     '8'),
    (2, 'Paciente.java',                                 '9'),
    (2, 'Validador.java',                               '10'),
    (2, 'ArchivoManager.java',                          '11'),
    (2, 'ExpedienteManager.java',                       '12'),
    (2, 'MenuUtils.java',                               '16'),
    (0, '3. Conclusión',                                '18'),
    (0, '4. Referencias',                               '19'),
]

def build_toc(doc):
    add_heading1(doc, 'Tabla de Contenido')
    doc.add_paragraph()

    tbl = doc.add_table(rows=len(TOC_ENTRIES), cols=2)
    tbl.style = 'Table Grid'
    # Remove table borders
    tbl_el = tbl._tbl
    tbl_pr = tbl_el.find(qn('w:tblPr'))
    if tbl_pr is None:
        tbl_pr = OxmlElement('w:tblPr')
        tbl_el.insert(0, tbl_pr)
    tbl_borders = OxmlElement('w:tblBorders')
    for side in ('top', 'bottom', 'left', 'right', 'insideH', 'insideV'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'none')
        tbl_borders.append(b)
    tbl_pr.append(tbl_borders)

    for i, (lvl, title, page) in enumerate(TOC_ENTRIES):
        row = tbl.rows[i]
        indent = lvl * 0.25
        cell_title = row.cells[0]
        cell_page  = row.cells[1]

        # Title cell
        cell_title.text = ''
        p = cell_title.paragraphs[0]
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.left_indent = Inches(indent)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(title)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.font.bold = (lvl == 0)

        # Page number cell
        cell_page.text = ''
        p2 = cell_page.paragraphs[0]
        p2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p2.paragraph_format.space_before = Pt(2)
        p2.paragraph_format.space_after = Pt(2)
        run2 = p2.add_run(page)
        run2.font.name = 'Times New Roman'
        run2.font.size = Pt(12)

        # Set column widths
        for col, w in [(0, 5.5), (1, 0.5)]:
            row.cells[col].width = Inches(w)

# ── Introducción ──────────────────────────────────────────────────────────────

def build_intro(doc):
    add_heading1(doc, '1. Introducción')
    for para_text in INTRO:
        add_body_para(doc, para_text, indent=True)

# ── Bitácora ──────────────────────────────────────────────────────────────────

def build_bitacora(doc):
    add_heading1(doc, '2. Diseño')
    add_heading2(doc, '2.1 Bitácora de Trabajo')
    add_body_para(doc, (
        'A continuación se detalla el registro de las sesiones de trabajo colaborativo '
        'realizadas durante el desarrollo del proyecto. Cada sesión refleja las tareas '
        'asignadas y los avances concretos obtenidos.'
    ), indent=False)

    doc.add_paragraph()

    tbl = doc.add_table(rows=1 + len(BITACORA), cols=5)
    tbl.style = 'Table Grid'

    headers = ['Fecha', 'Hora', 'Asistencia', 'Tarea Asignada', 'Avances Obtenidos']
    widths  = [0.8, 1.0, 1.1, 2.0, 2.1]
    for i, (hdr_txt, w) in enumerate(zip(headers, widths)):
        cell = tbl.rows[0].cells[i]
        set_cell_font(cell, hdr_txt, bold=True, size=11, bg='1A1A1A')
        cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        cell.width = Inches(w)

    for row_i, (fecha, hora, asist, tarea, avance) in enumerate(BITACORA):
        row = tbl.rows[row_i + 1]
        fill = 'F8F8F8' if row_i % 2 == 1 else 'FFFFFF'
        for col_i, text in enumerate([fecha, hora, asist, tarea, avance]):
            set_cell_font(row.cells[col_i], text, size=10, bg=fill)
            row.cells[col_i].width = Inches(widths[col_i])

# ── Flujogramas ───────────────────────────────────────────────────────────────

FLUJOGRAMA_DESCS = [
    ('2.2.1 Flujo Principal del Sistema — Main.java',
     'Muestra el flujo completo desde el INICIO hasta el FIN: bienvenida, login, '
     'bucle principal con switch de 6 opciones, y despedida al confirmar la salida.'),
    ('2.2.2 Proceso de Login — MenuUtils.realizarLogin()',
     'Ciclo MIENTRAS intentos < 3: solicita usuario y contraseña via JOptionPane, '
     'valida credenciales con arreglos paralelos; termina con return true o false.'),
    ('2.2.3 Crear Nuevo Expediente — ExpedienteManager.crearExpediente()',
     'Ciclos DO-WHILE para validar cada campo del expediente, genera número automático '
     'EXP-DDDDD, crea objeto Paciente y lo guarda en el archivo.'),
    ('2.2.4 Buscar por Número — ExpedienteManager.buscarPorNumero()',
     'Solicita número de expediente, recorre el arreglo Paciente[100] con FOR '
     'y .equals(), muestra el resultado o informa si no se encontró.'),
    ('2.2.5 Buscar por Fecha de Visita — ExpedienteManager.buscarPorFechaVisita()',
     'Valida formato de fecha con charAt()/length() en un bucle MIENTRAS, '
     'recorre todos los expedientes y muestra los que coinciden.'),
    ('2.2.6 Actualizar Expediente — ExpedienteManager.actualizarExpediente()',
     'Busca el expediente, entra en ciclo MIENTRAS continuar=true con submenú '
     'switch de 11 opciones para actualizar campos individuales.'),
    ('2.2.7 Listar Todos los Expedientes — ExpedienteManager.listarTodosLosExpedientes()',
     'Carga todos los expedientes, construye encabezado con FOR (separadores), '
     'recorre el arreglo y muestra tabla completa en JOptionPane.'),
]

def build_flujogramas(doc):
    add_heading2(doc, '2.2 Flujogramas')
    add_body_para(doc, (
        'Los flujogramas del sistema se presentan en el archivo adjunto '
        'Flujogramas_COMP2315.pdf (7 páginas, estilo Raptor). A continuación '
        'se describe el flujo de cada diagrama:'
    ), indent=False)

    for titulo, desc in FLUJOGRAMA_DESCS:
        add_heading3(doc, titulo)
        add_body_para(doc, desc, indent=False)
        # Placeholder para imagen
        p = doc.add_paragraph()
        set_para_format(p, align=WD_ALIGN_PARAGRAPH.CENTER,
                        space_before=4, space_after=8)
        run = p.add_run('[Ver Flujogramas_COMP2315.pdf — página correspondiente]')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        run.font.italic = True
        run.font.color.rgb = RGBColor(100, 100, 100)

# ── Menú y Submenú ────────────────────────────────────────────────────────────

def build_menus(doc):
    add_heading2(doc, '2.3 Menú Principal y Submenú de Actualización')
    add_body_para(doc, MENU_PRINCIPAL_TEXT, indent=True)

    tbl = doc.add_table(rows=1 + len(MENU_OPTIONS), cols=3)
    tbl.style = 'Table Grid'
    for i, hdr_txt in enumerate(['Opción', 'Descripción', 'Método invocado']):
        set_cell_font(tbl.rows[0].cells[i], hdr_txt, bold=True, size=11, bg='1A3A5C')
        tbl.rows[0].cells[i].paragraphs[0].paragraph_format.alignment = \
            WD_ALIGN_PARAGRAPH.CENTER
        tbl.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = \
            RGBColor(255, 255, 255)
    for row_i, (opt, desc, method) in enumerate(MENU_OPTIONS):
        row = tbl.rows[row_i + 1]
        fill = 'F8F8F8' if row_i % 2 == 1 else 'FFFFFF'
        for col_i, text in enumerate([opt, desc, method]):
            set_cell_font(row.cells[col_i], text, size=10, bg=fill)
        tbl.rows[row_i + 1].cells[0].paragraphs[0].paragraph_format.alignment = \
            WD_ALIGN_PARAGRAPH.CENTER

    for i, w in enumerate([0.5, 2.0, 3.5]):
        for row in tbl.rows:
            row.cells[i].width = Inches(w)

    doc.add_paragraph()
    add_body_para(doc, SUBMENU_TEXT, indent=True)

    tbl2 = doc.add_table(rows=1 + len(SUBMENU_OPTIONS), cols=2)
    tbl2.style = 'Table Grid'
    for i, hdr_txt in enumerate(['Opción', 'Campo a Actualizar']):
        set_cell_font(tbl2.rows[0].cells[i], hdr_txt, bold=True, size=11, bg='1A3A5C')
        tbl2.rows[0].cells[i].paragraphs[0].paragraph_format.alignment = \
            WD_ALIGN_PARAGRAPH.CENTER
        tbl2.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = \
            RGBColor(255, 255, 255)
    for row_i, (opt, campo) in enumerate(SUBMENU_OPTIONS):
        row = tbl2.rows[row_i + 1]
        fill = 'F8F8F8' if row_i % 2 == 1 else 'FFFFFF'
        for col_i, text in enumerate([opt, campo]):
            set_cell_font(row.cells[col_i], text, size=10, bg=fill)
        tbl2.rows[row_i + 1].cells[0].paragraphs[0].paragraph_format.alignment = \
            WD_ALIGN_PARAGRAPH.CENTER

    for i, w in enumerate([0.5, 5.5]):
        for row in tbl2.rows:
            row.cells[i].width = Inches(w)

# ── Código Fuente ─────────────────────────────────────────────────────────────

def build_codigo(doc):
    add_heading2(doc, '2.4 Código Fuente')
    add_body_para(doc, (
        'A continuación se presenta el código fuente completo de las seis clases '
        'que conforman el Sistema de Expedientes Médicos del Dr. Rodríguez, '
        'debidamente identificadas y comentadas en español.'
    ), indent=True)

    for fname, desc in JAVA_FILES:
        add_page_break(doc)
        add_heading3(doc, fname)
        add_body_para(doc, desc, indent=False, space_after=4)

        fpath = os.path.join(JAVA_DIR, fname)
        if os.path.exists(fpath):
            with open(fpath, encoding='utf-8') as f:
                code = f.read()
        else:
            code = f'[Archivo {fname} no encontrado]'

        # Split into chunks to avoid single massive run
        lines = code.split('\n')
        # Group lines into paragraphs of ~50 lines for performance
        chunk_size = 50
        for i in range(0, len(lines), chunk_size):
            chunk = '\n'.join(lines[i:i + chunk_size])
            add_monospace(doc, chunk, size=8)

# ── Conclusión ────────────────────────────────────────────────────────────────

def build_conclusion(doc):
    add_heading1(doc, '3. Conclusión')
    for para_text in CONCLUSION:
        add_body_para(doc, para_text, indent=True)

# ── Referencias ───────────────────────────────────────────────────────────────

def build_referencias(doc):
    add_heading1(doc, '4. Referencias')
    for ref in REFERENCIAS:
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        pf.first_line_indent = Inches(-0.5)
        pf.left_indent = Inches(0.5)
        pf.space_before = Pt(0)
        pf.space_after = Pt(6)
        run = p.add_run(ref)
        set_font(run)

# ── Main ──────────────────────────────────────────────────────────────────────

def main():
    doc = Document()

    # Default style
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    # Section margins (applies to all sections unless overridden)
    section = doc.sections[0]
    set_section_margins(section)

    # Page numbers in header
    add_page_number_header(doc)

    # ── Portada (page 1, no section break needed — header still shows) ──
    build_portada(doc)
    add_page_break(doc)

    # ── Tabla de Contenido (page 2) ──────────────────────────────────────
    build_toc(doc)
    add_page_break(doc)

    # ── Introducción (page 3) ────────────────────────────────────────────
    build_intro(doc)
    add_page_break(doc)

    # ── Bitácora (page 4) ────────────────────────────────────────────────
    build_bitacora(doc)
    add_page_break(doc)

    # ── Flujogramas (page 5–6) ───────────────────────────────────────────
    build_flujogramas(doc)
    add_page_break(doc)

    # ── Menú y Submenú (page 7) ──────────────────────────────────────────
    build_menus(doc)
    add_page_break(doc)

    # ── Código fuente (page 8+) ───────────────────────────────────────────
    build_codigo(doc)
    add_page_break(doc)

    # ── Conclusión ────────────────────────────────────────────────────────
    build_conclusion(doc)
    add_page_break(doc)

    # ── Referencias ───────────────────────────────────────────────────────
    build_referencias(doc)

    os.makedirs(os.path.dirname(OUT_PATH), exist_ok=True)
    doc.save(OUT_PATH)
    print(f'Documento generado: {OUT_PATH}')


if __name__ == '__main__':
    main()
